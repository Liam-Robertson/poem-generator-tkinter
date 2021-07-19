import tkinter as tk
from tkinter import ttk
import math
import docx
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import os
import datetime 
from operator import itemgetter
import sys
from termcolor import colored
importFilePath = os.path.join(os.path.dirname(os.path.realpath(__file__)), 'utils')
sys.path.append(importFilePath)
import customUtils as cu
from customUtils import sysStatus, debugVar, convertPath
import poemFinder as pf

poemListPath = os.path.join(os.path.dirname(os.path.realpath(__file__)), 'syllabaryPoems')
remainingPoemList = [filename[:-4] for filename in os.listdir(poemListPath) if (os.path.isfile(os.path.join(poemListPath, filename)) and filename[-4:] == '.xml')]
poemListSplit = [[int(poemName.split("-")[0]), int(poemName.split("-")[1]), int(poemName.split("-")[2])] for poemName in remainingPoemList]
poemListSplit.sort()
xUniqueCoordList = sorted(set([poemNameSplit[0] for poemNameSplit in poemListSplit]))
startingPoem = '1-1-1'
numOfPoems = 1
poemOrder = 'forwards'
runPoemFinderVar = False

class PoetryGenerator(tk.Tk):
	def __init__(self, *args, **kwargs):
		tk.Tk.__init__(self, *args, **kwargs)

		# Initalise master window container (1x1 grid that expands in any direction) 
		self.grid_rowconfigure(0, weight=1)
		self.grid_columnconfigure(0, weight=1)

		# Initialise secondary container (3x1 grid, all rows expand horizontally, top row can expand vertically also)
		main_container = tk.Frame(self, padx=30, pady=30, bg='blanched almond')
		main_container.grid(column=0, row=0, sticky = "nsew")
		main_container.grid_rowconfigure(0, weight = 1)
		main_container.grid_columnconfigure(0, weight = 1)

		# Get pageOne as an object by passing through its parameters
		# Then switch master window container to pageOne
		self.frames = {}
		for args in (PageOne, PageConfirmation, PagePoemList, LoadingPage):
			frame = args(main_container, self)
			self.frames[args] = frame
			frame.grid(row = 0, column = 0, sticky = "nsew")
		self.show_frame(PageOne)

	# Method to switch master window container to a new window object
	def show_frame(self, pointer):
		frame = self.frames[pointer]
		frame.tkraise()
	
class PageOne(tk.Frame):
	def __init__(self, parent, controller):
		tk.Frame.__init__(self, parent)  
		self.controller = controller 
		self.parent = parent
		self.poemListPath = os.path.join(os.path.dirname(os.path.realpath(__file__)), 'syllabaryPoems')
		self.poemFilenameList = [filename for filename in os.listdir(self.poemListPath) if (os.path.isfile(os.path.join(self.poemListPath, filename)) and filename[-4:] == '.xml')]
		self.unfilteredPoemFilenameList = [filename for filename in os.listdir(self.poemListPath)]  
		self.numOfPoemsWarning = tk.StringVar()
		self.startingPoemWarning = tk.StringVar()
		self.poemOrderWarning = tk.StringVar()

		# Initalise master window container (1x1 grid that expands in any direction) 
		self.grid_rowconfigure(0, weight=1)
		self.grid_columnconfigure(0, weight=1)

		# Initialise secondary container (3x1 grid, all rows expand horizontally, top row can expand vertically also)
		mainSubframe = tk.Frame(self, bg='navajo white', padx=100)
		mainSubframe.grid(column=0, row=0, sticky = "nsew")
		mainSubframe.grid_rowconfigure(0, weight = 1)
		mainSubframe.grid_columnconfigure(0, weight = 1)

		# Initialise pageOne container (4x1 grid where all rows expand horizontally and row 1 expands vertically as well)
		mainSubframe.columnconfigure(0, weight = 1)
		mainSubframe.rowconfigure(0, weight = 1)
		mainSubframe.rowconfigure(1, weight = 1)
		mainSubframe.rowconfigure(2, weight = 1)
		mainSubframe.rowconfigure(3, weight = 1)
		mainSubframe.rowconfigure(4, weight = 1)

		mainSubframe.grid_rowconfigure(0, weight=1)
		mainSubframe.grid_columnconfigure(0, weight=1)

		# Create subcontainer 1 - number of poems
		subframe1 = tk.Frame(mainSubframe, bg='gray99',  relief=tk.RIDGE,  bd=5)
		subframe1.grid_columnconfigure(0, weight=1)
		subframe1.grid_rowconfigure(0, weight=1)
		subframe1.grid_rowconfigure(1, weight=1)
		subframe1.grid_rowconfigure(2, weight=1)
		subframe1.grid(row = 1, sticky = "nsew", pady=10)

		# Create subcontainer 2 - starting poem
		subframe2 = tk.Frame(mainSubframe, bg='gray99',  relief=tk.RIDGE,  bd=5)
		subframe2.grid_columnconfigure(0, weight=1)
		subframe2.grid_rowconfigure(0, weight=1)
		subframe2.grid_rowconfigure(1, weight=1)
		subframe2.grid_rowconfigure(2, weight=1)
		subframe2.grid_rowconfigure(3, weight=1)
		subframe2.grid_rowconfigure(4, weight=1)
		subframe2.grid(row = 2, sticky = "nsew", pady=10)

		# Create subcontainer 3 - radio buttons
		radioMainContainer = tk.Frame(mainSubframe, bg='gray99',  relief=tk.RIDGE,  bd=5)
		radioMainContainer.grid_columnconfigure(0, weight=1)
		radioMainContainer.grid_rowconfigure(0, weight=1)
		radioMainContainer.grid(row = 3, sticky = "nsew", pady=10)

		# Create subcontainer 3 - radio buttons
		subframe3 = tk.Frame(radioMainContainer, bg='gray99')
		subframe3.grid_columnconfigure(0, weight=1)
		subframe3.grid_columnconfigure(1, weight=1)
		subframe3.grid_rowconfigure(0, weight=1)
		subframe3.grid_rowconfigure(1, weight=1)
		subframe3.grid_rowconfigure(2, weight=1)
		subframe3.grid(row = 0)

		# Create subcontainer 4 - submit button
		subframe4 = tk.Frame(mainSubframe, bg='gray99')
		subframe4.grid_columnconfigure(0, weight=1)
		subframe4.grid_rowconfigure(0, weight=1)
		subframe4.grid(row = 4, pady=20)

		# Create subcontainer for radio buttons
		radioSubframe = tk.Frame(subframe3, bg='gray99')
		radioSubframe.grid_columnconfigure(0, weight=1)
		radioSubframe.grid_columnconfigure(1, weight=1)
		radioSubframe.grid_rowconfigure(0, weight=1)
		radioSubframe.grid(row = 1)

		# Initialising styling variables
		pageOneStyle = ttk.Style(master=self)
		pageOneStyle.configure('W.TLabel', font=('calibre', 20, 'normal'), background='gray99')
		pageOneStyle.configure('W.TEntry', font=('calibre', 20, 'normal'), background='gray99', relief=tk.FLAT, borderwidth=15)
		pageOneStyle.configure('W.TRadiobutton', font=('calibre', 20, 'normal'), background='gray99')
		pageOneStyle.configure('W.TButton', font=('calibre', 20, 'normal'), height=100)

		# Initialising variables
		numOfPoemsTk = tk.IntVar()
		startingPoemTk = tk.StringVar()
		poemOrderTk = tk.StringVar()

		# Creating a title and assigning it to the first row of the window 
		titleLabel = ttk.Label(mainSubframe, text = "Syllabary Poem Generator",  font=("Times New Roman", 40, 'italic'), background='blanched almond')
		titleLabel.grid(row = 0, pady=(10, 0), ipady=0, sticky = "nsew")
		titleLabel.configure(anchor='center')

		# Creating user prompt for getting number of poems
		titleLabel = ttk.Label(subframe1, text = "How many poems do you want to generate?", font=("calibre", 20, 'normal'), style='W.TLabel')
		titleLabel.grid(row = 0, pady=(5, 0))

		# User entry bar for getting number of poems
		poemNumEntry = ttk.Entry(subframe1, textvariable=numOfPoemsTk, font=('calibre', 20, 'normal'))
		poemNumEntry.grid(row = 1, pady=(15, 0))

		# Creating user output - number of poems
		numPoemsOutputLabel = tk.Label(subframe1, textvariable=self.numOfPoemsWarning, background='gray99', foreground='red', font=('calibre', 13, 'normal'))
		numPoemsOutputLabel.grid(row = 2)

		# Creating user prompt for getting number of poems
		titleLabel = ttk.Label(subframe2, text = "What poem do you want to select?", font=("calibre", 20, 'normal'), style='W.TLabel')
		titleLabel.grid(row = 0, pady=(5, 0))

		# Creating user prompt for getting number of poems
		titleLabel = ttk.Label(subframe2, text = "Enter poem in format x-x-x (For example: 2-10-15)", font=("calibre", 12, 'italic'), background='gray99', foreground='gray15')
		titleLabel.grid(row = 1, pady=(0, 0))

		# User entry bar for getting poem order
		startingPoemEntry = ttk.Entry(subframe2, textvariable=startingPoemTk, style='W.TEntry', font=('calibre', 20, 'normal'))
		startingPoemEntry.grid(row = 2, pady=(15, 0))

		# Creating submit button
		poemListButton = ttk.Button(subframe2, text="View Poem List", command=lambda: self.showPoemList(PagePoemList), width=20, style='W.TButton') 
		poemListButton.grid(row = 3,  pady=(10, 0))

		# Creating user output - starting poem
		startPoemOutputLabel = tk.Label(subframe2, textvariable=self.startingPoemWarning, background='gray99', foreground='red', font=('calibre', 13, 'normal'))
		startPoemOutputLabel.grid(row = 4)

		# Creating user prompt for getting poem order
		titleLabel = ttk.Label(subframe3, text = "Do you want the output poem list to start or end with the selected poem?", font=("calibre", 20, 'normal'), style='W.TLabel')
		titleLabel.grid(row = 0, pady=(5, 0))

		# Creating checkboxes to let the user input poem order
		radioButton1 = ttk.Radiobutton(radioSubframe, text="Start", variable=poemOrderTk, value="starting", style='W.TRadiobutton')
		radioButton2 = ttk.Radiobutton(radioSubframe, text="End", variable=poemOrderTk, value="ending", style='W.TRadiobutton')
		radioButton1.grid(row = 1, column=0, padx=10, pady=(5, 0))
		radioButton2.grid(row = 1, column=1, padx=10, pady=(5, 0))    

		# Creating user output - radio buttons
		poemOrderOutputLabel = tk.Label(subframe3, textvariable=self.poemOrderWarning, background='gray99', foreground='red', font=('calibre', 13, 'normal'))
		poemOrderOutputLabel.grid(row = 2)

		# Creating submit button
		submitButton = ttk.Button(subframe4, text="Submit", command=lambda: self.submit(LoadingPage, numPoemsOutputLabel, startPoemOutputLabel, poemOrderOutputLabel, startingPoemTk, numOfPoemsTk, poemOrderTk), width=30, style='W.TButton') 
		submitButton.grid(row = 0)

	def submit(self, pointer, numPoemsOutputLabel, startPoemOutputLabel, poemOrderOutputLabel, startingPoemTk, numOfPoemsTk, poemOrderTk):
		global startingPoem
		global numOfPoems
		global poemOrder
		global runPoemFinderVar
		startingPoem = startingPoemTk.get()
		numOfPoems = numOfPoemsTk.get()
		poemOrder = poemOrderTk.get()
		print('submit pressure')

		self.checkInputs(numPoemsOutputLabel)
		numPoemsOutputLabel.config(text=self.numOfPoemsWarning.get())
		startPoemOutputLabel.config(text=self.startingPoemWarning.get())
		poemOrderOutputLabel.config(text=self.poemOrderWarning.get()) 
		if self.numOfPoemsWarning.get() == '' and self.startingPoemWarning.get() == '' and self.poemOrderWarning.get() == '':
			runPoemFinderVar = True
			self.controller.show_frame(pointer)
			LoadingPage(self.parent, self.controller)

	def showPoemList(self, pointer):
		self.controller.show_frame(pointer)

	def checkInputs(self, numPoemsOutputLabel):
		# poemList check
		numDiscardedFiles = len(self.unfilteredPoemFilenameList) - len(self.poemFilenameList)
		if len(self.poemFilenameList) > 0: 
			print(f'Success: poemList directory is valid. {len(self.poemFilenameList)} files have been found.')
			if len(self.poemListPath) != len(self.unfilteredPoemFilenameList):
				print(f'Warning: {numDiscardedFiles} files have been discarded as they were not in a valid format')
		else:
			print('Error: poem directory is empty. Please select a valid input directory')

		# numOfPoems check
		if numOfPoems > 0:
			if isinstance(numOfPoems, int):
				if numOfPoems <= len(self.poemFilenameList):
					print('Success: number of poems is valid')
					self.numOfPoemsWarning.set('')
				else:
					self.numOfPoemsWarning.set(f'Error: there are not enough available poems to fulfil that request. Please select a number less than {len(self.poemFilenameList) + 1}')
			else:
				self.numOfPoemsWarning.set('Error: number of poems is not integer. Please insert an integer')
		else:
			self.numOfPoemsWarning.set('Error: number of poems must be a positive number')
			
		# startingPoem check
		startingPoemSplit = startingPoem.split('-')
		if len(startingPoemSplit) == 3 and '' not in startingPoemSplit:
			if startingPoemSplit[0].isnumeric() and startingPoemSplit[1].isnumeric() and startingPoemSplit[2].isnumeric():
				if f'{startingPoem}.xml' in self.poemFilenameList:
					print('Success: Starting poem is valid')
					self.startingPoemWarning.set('')
				else:
					self.startingPoemWarning.set('Error: poem not in list. Please select a different poem. Click the poem list button for information about available poems')
			else:
				self.startingPoemWarning.set('Error: format is correct but not all the characters were numbers.  Please enter poem in the format: "x-x-x" where x is a number')
		else:
			self.startingPoemWarning.set('Error: poem was not entered in the correct format. Please enter poem in the format: "x-x-x" where x is a number')

		# Poem order check
		if poemOrder == 'starting' or poemOrder == 'ending':
			print('Success: poem order is valid')
			self.poemOrderWarning.set('')
		elif poemOrder == '':
			self.poemOrderWarning.set('Error: output empty. Select "Start" or "End"')
	
class LoadingPage(tk.Frame):
	def __init__(self, parent, controller):
		tk.Frame.__init__(self, parent)    
		self.controller = controller

		# Initialising styling variables
		loadingStyle = ttk.Style(master=self)
		loadingStyle.configure('L.TLabel', font=('calibre', 20, 'normal'), background='gray99')
		loadingStyle.configure('Lloading.TLabel', font=('calibre', 20, 'italic'), background='gray99')
		loadingStyle.configure('L.TEntry', font=('calibre', 20, 'normal'), background='gray99', relief=tk.FLAT, borderwidth=15)
		loadingStyle.configure('L.TRadiobutton', font=('calibre', 20, 'normal'), background='gray99')
		loadingStyle.configure('L.TButton', font=('calibre', 20, 'normal'), height=100)   
		loadingStyle.configure('L.TFrame', background='gray99',  relief=tk.RIDGE,  bd=15) 
		loadingStyle.configure('LMain.TFrame', background='navajo white') 

		# Initalise master window container (1x1 grid that expands in any direction) 
		self.grid_rowconfigure(0, weight=1)
		self.grid_columnconfigure(0, weight=1)

		# Initialise secondary container (3x1 grid, all rows expand horizontally, top row can expand vertically also)
		mainSubframe = ttk.Frame(self, style='LMain.TFrame')
		mainSubframe.grid(column=0, row=0, sticky = "nsew")
		mainSubframe.grid_rowconfigure(0, weight = 1)
		mainSubframe.grid_columnconfigure(0, weight = 1)

		# Create subcontainer 1 - number of poems
		allItemsSubframe = ttk.Frame(mainSubframe,  style='L.TFrame')
		allItemsSubframe.grid_columnconfigure(0, weight=1)
		allItemsSubframe.grid_rowconfigure(0, weight=1)
		allItemsSubframe.grid_rowconfigure(1, weight=1)
		allItemsSubframe.grid_rowconfigure(2, weight=1)
		allItemsSubframe.grid_rowconfigure(3, weight=1)
		allItemsSubframe.grid(column=0, row=0, ipadx=20, ipady=20)

		selectingPoemsLabel = ttk.Label(allItemsSubframe, text='Selecting poems...', style='Lloading.TLabel')
		selectingPoemsLabel.grid(column=0, row=0, pady=(30, 0))

		WritingFileLabel = ttk.Label(allItemsSubframe, text='Writing poems to file...', style='Lloading.TLabel')
		WritingFileLabel.grid(column=0, row=2, pady=(20, 0))

		self.update_idletasks()

		global runPoemFinderVar
		print(f"run poem finder {runPoemFinderVar}")
		outputPoemDict = self.runPoemFinder(allItemsSubframe)
		self.creatingTextDocumentOutput(outputPoemDict, startingPoem, numOfPoems, allItemsSubframe)

	def runPoemFinder(self, parentFrame):
		outputPoemList = pf.createOutputPoemList(startingPoem, numOfPoems, poemOrder, parentFrame)
		outputPoemDict = pf.readingXML(outputPoemList)
		return outputPoemDict
		#pf.creatingTextDocumentOutput(root, outputPoemDict, startingPoem, numOfPoems, parentFrame)

	def createProgressBar(self, parentFrame, rowVar):
		selectingPoemsProgress = ttk.Progressbar(parentFrame, orient=tk.HORIZONTAL, length=300, mode='determinate')
		selectingPoemsProgress.grid(column=0, row=rowVar, padx=30)
		return selectingPoemsProgress

	def creatingTextDocumentOutput(self, poemDict, startingPoem, numOfPoems, parentFrame):
		"""Outputting to a word document"""
		doc = docx.Document()
		stepSize = 10
		writeProgressBar = self.createProgressBar(parentFrame, 3)
		for poemName in poemDict.keys():
			style = doc.styles['Normal']
			font = style.font
			font.name = 'Helvetica'
			font.size = Pt(14)
			par = doc.add_paragraph()
			par.paragraph_format.line_spacing = 1
			run1 = par.add_run(poemName)
			run1.add_break(WD_BREAK.LINE)
			run1.add_break(WD_BREAK.LINE)
			if poemDict[poemName]['Title'] != None:
				if poemDict[poemName]['Title'].isspace() == False:
					run2 = par.add_run(poemDict[poemName]['Title'])
					run2.bold = True
					run2.add_break(WD_BREAK.LINE)
					run2.add_break(WD_BREAK.LINE)
			run3 = par.add_run(poemDict[poemName]['Content'])
			run3.add_break(WD_BREAK.PAGE)
			doc.save(f'PoemGenerator_{startingPoem}_{numOfPoems}.docx')
			writeProgressBar['value'] += stepSize
			print(writeProgressBar['value'])
			self.update_idletasks
		print('finished writing')


class PageConfirmation(tk.Frame):
	def __init__(self, parent, controller):
		tk.Frame.__init__(self, parent)    
		self.controller = controller     

		# Initialising styling variables
		pageOneStyle = ttk.Style(master=self)
		pageOneStyle.configure('C.TLabel', font=('calibre', 20, 'normal'), background='gray99')
		pageOneStyle.configure('C2.TLabel', font=('calibre', 30, 'italic'), background='gray99')
		pageOneStyle.configure('C.TEntry', font=('calibre', 20, 'normal'), background='gray99', relief=tk.FLAT, borderwidth=15)
		pageOneStyle.configure('C.TRadiobutton', font=('calibre', 20, 'normal'), background='gray99')
		pageOneStyle.configure('C.TButton', font=('calibre', 20, 'normal'), height=100)   

		# Initalise master window container (1x1 grid that expands in any direction) 
		self.grid_rowconfigure(0, weight=1)
		self.grid_columnconfigure(0, weight=1)

		# Initialise secondary container (3x1 grid, all rows expand horizontally, top row can expand vertically also)
		mainSubframe = tk.Frame(self, bg='navajo white')
		mainSubframe.grid(column=0, row=0, sticky = "nsew")
		mainSubframe.grid_rowconfigure(0, weight = 1)
		mainSubframe.grid_columnconfigure(0, weight = 1)

		# Create subcontainer 1 - number of poems
		allItemsSubframe = tk.Frame(mainSubframe, bg='navajo white')
		allItemsSubframe.grid_columnconfigure(0, weight=1)
		allItemsSubframe.grid_rowconfigure(0, weight=1)
		allItemsSubframe.grid_rowconfigure(1, weight=1)
		allItemsSubframe.grid(column=0, row=0, ipady=15, ipadx=30)

		# Create subcontainer 1 - number of poems
		allText = tk.Frame(allItemsSubframe, bg='gray99',  relief=tk.RIDGE,  bd=5)
		allText.grid_columnconfigure(0, weight=1)
		allText.grid_rowconfigure(0, weight=1)#
		allText.grid_rowconfigure(1, weight=1)
		allText.grid(column=0, row=0, ipady=15, ipadx=30, pady=(0, 40))

		successLabel = ttk.Label(allText, text='Success!', style='C2.TLabel')
		successInfoLabel = ttk.Label(allText, text='Poem document is in Downloads folder', style='C.TLabel')
		successLabel.grid(column=0, row=0, pady=(0, 10))
		successInfoLabel.grid(column=0, row=1, pady=0)

		# Creating submit button
		backButton = ttk.Button(allItemsSubframe, text="Back", command=lambda: self.backButton(PageOne), width=30, style='W.TButton') 
		backButton.grid(row = 1)

	def backButton(self, pointer):
		self.controller.show_frame(pointer)

class PagePoemList(tk.Frame):
	def __init__(self, parent, controller):
		tk.Frame.__init__(self, parent)      
		self.controller = controller    
		self.parent = parent    

		# Initialising styling variables
		PagePoemList = ttk.Style(master=self)
		PagePoemList.configure('W.TLabel', font=('calibre', 20, 'normal'), background='gray99')
		PagePoemList.configure('W.TEntry', font=('calibre', 20, 'normal'), background='gray99', relief=tk.FLAT, borderwidth=15)
		PagePoemList.configure('W.TRadiobutton', font=('calibre', 20, 'normal'), background='gray99')
		PagePoemList.configure('W.TButton', font=('calibre', 20, 'normal'), height=100)
		PagePoemList.configure('W.TNotebook', font=('calibre', 20, 'normal'))

		# Initalise master window container (1x1 grid that expands in any direction) 
		self.grid_rowconfigure(0, weight=1)
		self.grid_columnconfigure(0, weight=1)

		# Initialise main container 
		mainSubframe = tk.Frame(self, bg='navajo white')
		mainSubframe.grid(column=0, row=0, sticky = "nsew")
		mainSubframe.grid_rowconfigure(0, weight = 1)
		mainSubframe.grid_columnconfigure(0, weight = 1)

		# Initialise container for tabs and back button
		subframe = tk.Frame(mainSubframe)
		subframe.grid(column=0, row=0, sticky = "nsew")
		subframe.grid_rowconfigure(0, weight = 1)
		subframe.grid_rowconfigure(1, weight = 1)
		subframe.grid_columnconfigure(0, weight = 1)

		# Initialising tabs
		poemListNotebook = ttk.Notebook(subframe, style='W.TNotebook')
		poemListNotebook.grid(column=0, row=0, sticky = "nsew")
		poemListNotebook.grid_rowconfigure(0, weight = 1)
		poemListNotebook.grid_rowconfigure(1, weight = 1)
		poemListNotebook.grid_columnconfigure(0, weight = 1)

		# Initialise container for back button
		buttonSubframe = tk.Frame(subframe, bg='blanched almond')
		buttonSubframe.grid_rowconfigure(0, weight = 1)
		buttonSubframe.grid_columnconfigure(0, weight = 1)
		buttonSubframe.grid(column=0, row=1, sticky = "nsew")

		# # Initialise container for text
		xCoordColumnNum = 13
		tabNameDict = {}
		for currentXCoord in xUniqueCoordList:
			xCoordListSplit = [poemNameSplit for poemNameSplit in poemListSplit if poemNameSplit[0] == currentXCoord]
			xCoordListStr = [[str(poemNameSplit[0]), str(poemNameSplit[1]), str(poemNameSplit[2])] for poemNameSplit in xCoordListSplit]
			xCoordList = ['-'.join(poemNameSplit) for poemNameSplit in xCoordListStr]
			xCoordRowNum = int(math.ceil(len(xCoordListSplit) / xCoordColumnNum))
	
			tabNameDict[currentXCoord] = f'tab{currentXCoord}Subrame'
			tabNameDict[currentXCoord] = tk.Frame(poemListNotebook, bg='navajo white')
			tabNameDict[currentXCoord].grid(column=0, row=0, sticky = "nsew")
			for i in range(xCoordRowNum):
				for j in range(xCoordColumnNum):
					tabNameDict[currentXCoord].grid_rowconfigure(i, weight=1)
					tabNameDict[currentXCoord].grid_rowconfigure(j, weight=1)
					titleCodeLabel = tk.Label(tabNameDict[currentXCoord], text=xCoordList[i+j], font=("arial", 15, 'normal'),  relief=tk.RIDGE, bg='gray99', bd=3)
					titleCodeLabel.grid(row=i, column=j, ipadx=3, ipady=0, padx=3, pady=3, sticky = "nsew")

		# Adding text to notebook
		for currentXCoord in xUniqueCoordList:
			poemListNotebook.add(tabNameDict[currentXCoord], text=f'page {currentXCoord}')

		# Creating back button
		backButton = ttk.Button(buttonSubframe, text="Back", command=lambda: self.backButton(PageOne), width=20, style='W.TButton') 
		backButton.grid(row = 2)

	def backButton(self, pointer):
		self.controller.show_frame(pointer)

def attempt2(frame):
	LoadingPage(tk.Frame)

def main():
	syllabaryGenerator = PoetryGenerator()

	syllabaryGenerator.geometry("1280x720")
	syllabaryGenerator.title("Syllabary Poem Generator")

	syllabaryGenerator.mainloop()

	
if __name__ == "__main__":
	#cu.preRunCleanUp()
	main()
	
	


