import xml.etree.ElementTree as ET
from lxml import etree
import os
import docx
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import itertools
from os.path import isfile, join
import requests
import datetime
import tkinter as tk
from tkinter import ttk 
import math
from termcolor import colored
import re 
from pathlib import Path
import sys
importFilePath1 = os.path.join(os.path.dirname(os.path.realpath(__file__)), 'utils')
sys.path.append(importFilePath1)
import customUtils as cu
from customUtils import sysStatus, debugVar, convertPath
poemListPath = os.path.join(os.path.dirname(os.path.realpath(__file__)), 'syllabaryPoems')
selectingProgressDict = {}
selectingProgressDict['Progress Value'] = [0]

def createProgressBar(parentFrame, rowVar):
	selectingPoemsProgress = ttk.Progressbar(parentFrame, orient=tk.HORIZONTAL, length=300, mode='determinate')
	selectingPoemsProgress.grid(column=0, row=rowVar, padx=30)
	return selectingPoemsProgress

def createOutputPoemList(startingPoem, numOfPoems, poemOrder, parentFrame):
	"""Creates a set for X,Y,Z coordinates where filename is X-Y-Z.xml
	Loops through (X+i)-Y-Z.xml until it finds a file which exists in the input directory 
	Once that is found it will loop through X-(Y+i)-Z.xml and so on """

	"""
	Notes: 
	- The remaining list needs to be in integer form to be sorted
	- itertools.cycle can't be used because you are removing items from the cycle 
	"""

	counter = 0
	stepSize = float(300 / numOfPoems)
	outputPoemList = []
	currentCoordinateAxis = 'X'
	coordinateAxisList = ['X', 'Y', 'Z']
	startingPoemSplit = [int(startingPoem.split("-")[0]), int(startingPoem.split("-")[1]), int(startingPoem.split("-")[2])]
	remainingPoemList = [filename[:-4] for filename in os.listdir(poemListPath) if (os.path.isfile(os.path.join(poemListPath, filename)) and filename[-4:] == '.xml')]
	remainingPoemListSplit = [[int(poemName.split("-")[0]), int(poemName.split("-")[1]), int(poemName.split("-")[2])] for poemName in remainingPoemList]
	remainingPoemListSplit.sort()
	startingPoemIndex = remainingPoemListSplit.index(startingPoemSplit)
	remainingPoemListSplit = remainingPoemListSplit[startingPoemIndex:] + remainingPoemListSplit[:startingPoemIndex]
	nextCoordAxisDict = {'X': 'Y', 'Y': 'Z', 'Z': 'X'}
	selectingPoemsCounter = [0]
	global selectingProgressDict
	selectingProgressDict["Progress Value"] = selectingPoemsCounter
	selectedProgressBar = createProgressBar(parentFrame, rowVar=1)

	availableCoordinatesDict = {}
	for index in range(len(coordinateAxisList)):
		orderedUniqueCoordinateList = sorted(set([poemNameSplit[index] for poemNameSplit in remainingPoemListSplit]))
		availableCoordinatesDict[coordinateAxisList[index]] = orderedUniqueCoordinateList
	targetCoordinate = (availableCoordinatesDict['X'])[0]

	while len(outputPoemList) < numOfPoems:
		for currentPoemSplit in remainingPoemListSplit:
			if len(outputPoemList) == numOfPoems:
				print(f'Success: length of output list is {len(outputPoemList)}')
				return outputPoemList
				break
			elif len(outputPoemList) != numOfPoems:
				# Getting initial coordinate axis
				currentAxisIndex = coordinateAxisList.index(currentCoordinateAxis)
				currentCoordinate = currentPoemSplit[currentAxisIndex]
				# print(f'current poem: {currentPoemSplit}')
				# print(f'current coordinate: {currentCoordinate}')
				# print(f'current axis: {currentCoordinateAxis}')
				# print(f'target coordinate: {targetCoordinate}')
				if currentCoordinate == targetCoordinate:
					# Moving to next coordinate axis					
					currentCoordinateAxis = nextCoordAxisDict[currentCoordinateAxis]
					currentAxisIndex = coordinateAxisList.index(currentCoordinateAxis)
					currentCoordinate = currentPoemSplit[currentAxisIndex]

					# Getting variables needed to create target coordinate
					originalLengthAvailableCoords = len(availableCoordinatesDict[currentCoordinateAxis])
					currentSetIndex = availableCoordinatesDict[currentCoordinateAxis].index(currentCoordinate)
					if currentSetIndex == len(availableCoordinatesDict[currentCoordinateAxis]) - 1:
						targetIndex = 0
					else:
						targetIndex = currentSetIndex + 1

					# Appending to success list, removing from current list and updating available coordinates
					outputPoemList.append(currentPoemSplit)
					remainingPoemListSplit = [poemNameSplit for poemNameSplit in remainingPoemListSplit if poemNameSplit != currentPoemSplit]
					availableCoordinatesDict[currentCoordinateAxis] = sorted(set([poemNameSplit[currentAxisIndex] for poemNameSplit in remainingPoemListSplit]))

					# Taking into account whether the set changed size
					newLengthAvailableCoords = len(availableCoordinatesDict[currentCoordinateAxis])
					if originalLengthAvailableCoords ==  newLengthAvailableCoords + 1:
						targetIndex -= 1

					# Creating target coordinate
					targetCoordinate = (availableCoordinatesDict[currentCoordinateAxis])[targetIndex]

					# Updating Progress bar
					selectingProgressDict['Progress Value'][0] += stepSize
					selectedProgressBar['value'] = selectingProgressDict['Progress Value'][0]
					print('select success')
					
def readingXML(outputPoemList):
	"""Creating a dictionary in the form (filename: [title, text])"""
	poemDict = {}	
	outputPoemList = [[str(poemName[0]), str(poemName[1]), str(poemName[2])] for poemName in outputPoemList]
	outputPoemList = [f"{'-'.join(poemName)}.xml" for poemName in outputPoemList]
	for poemName in outputPoemList:
		poemDict[poemName] = {}
		poemFilePath = os.path.join(poemListPath, poemName)
		try:
			tree = ET.parse(poemFilePath)
		except:
			with open(poemFilePath, 'r') as infile:
				data = infile.read()
				with open(poemFilePath, 'w') as outfile:
					data.encode('UTF-8')
					data = data.replace('&', '&amp;').replace('"', '&quot;').replace('“', '&quot;').replace('’', '&apos;').replace("'", "&apos;")
					outfile.write(data)
			tree = ET.parse(poemFilePath)
		root = tree.getroot()
		poemDict[poemName]['Title'] = (root[0].text)
		poemDict[poemName]['Content'] = root[2].text
	return poemDict

def creatingTextDocumentOutput(root, poemDict, startingPoem, numOfPoems, parentFrame):
	"""Outputting to a word document"""
	doc = docx.Document()
	stepSize = 10
	writeProgressBar = createProgressBar(parentFrame, 3)
	for poemName in poemDict.keys():
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Helvetica'
		font.size = Pt(14)
		par = doc.add_paragraph()
		par.paragraph_format.line_spacing = 1
		run1 = par.add_run(poemName)
		run1.add_break(WD_BREAK.LINE)
		run1.add_break(WD_BREAK.LINE)
		if poemDict[poemName]['Title'] != None:
			if poemDict[poemName]['Title'].isspace() == False:
				run2 = par.add_run(poemDict[poemName]['Title'])
				run2.bold = True
				run2.add_break(WD_BREAK.LINE)
				run2.add_break(WD_BREAK.LINE)
		run3 = par.add_run(poemDict[poemName]['Content'])
		run3.add_break(WD_BREAK.PAGE)
		doc.save(f'PoemGenerator_{startingPoem}_{numOfPoems}.docx')
		writeProgressBar['value'] += stepSize
		print(writeProgressBar['value'])
		root.update_idletasks
	print('finished writing')
		